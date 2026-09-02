"""Reading documents, translating them, and writing a new file beside them.

Two rules run through all of this:

  1. The original is only ever read. Output is always a new file, named with
     the target language code, and an existing file is never overwritten.
  2. Structure is worth more than inline formatting. Headings stay headings,
     lists stay lists, tables stay tables, subtitle timings stay exact. Bold on
     a single word inside a sentence does not survive, and cannot: translation
     reorders words, so there is no honest way to map a span back.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED = {".docx", ".txt", ".md", ".srt", ".vtt", ".pdf"}


@dataclass
class Job:
    """One file to translate, and what happened to it."""
    source: Path
    target_code: str
    output: Path | None = None
    blocks: int = 0
    translated: int = 0
    note: str = ""
    warnings: list[str] = field(default_factory=list)


def output_path(source: Path, target_code: str, suffix: str | None = None) -> Path:
    """A new name beside the original. Never returns an existing path."""
    ext = suffix or source.suffix
    stem = f"{source.stem}.{target_code}"
    candidate = source.with_name(stem + ext)
    n = 2
    while candidate.exists():
        candidate = source.with_name(f"{stem} ({n}){ext}")
        n += 1
    return candidate


# ---------------------------------------------------------------- plain text
def _translate_plain(job: Job, translate, progress=None) -> None:
    text = job.source.read_text(encoding="utf-8", errors="replace")
    # Blank lines separate blocks; keeping them keeps markdown and paragraphs.
    parts = re.split(r"(\n\s*\n)", text)
    body = [p for p in parts if not re.fullmatch(r"\n\s*\n", p)]
    job.blocks = sum(1 for p in body if p.strip())

    out, i = [], 0
    for p in parts:
        if re.fullmatch(r"\n\s*\n", p) or not p.strip():
            out.append(p)
            continue
        out.append(translate(p))
        i += 1
        job.translated = i
        if progress:
            progress(i, job.blocks)

    job.output = output_path(job.source, job.target_code)
    job.output.write_text("".join(out), encoding="utf-8")


# ------------------------------------------------------------------ subtitles
_SRT_TIME = re.compile(r"^\d{1,2}:\d{2}:\d{2}[,.]\d{1,3}\s*-->")


def _translate_subtitles(job: Job, translate, progress=None) -> None:
    """Translate the text lines only. Timings and cue numbers are untouched."""
    lines = job.source.read_text(encoding="utf-8-sig", errors="replace").splitlines()
    text_line_indexes = [
        i for i, ln in enumerate(lines)
        if ln.strip()
        and not _SRT_TIME.match(ln.strip())
        and not ln.strip().isdigit()
        and not ln.strip().upper().startswith("WEBVTT")
    ]
    job.blocks = len(text_line_indexes)

    for n, i in enumerate(text_line_indexes, start=1):
        lines[i] = translate(lines[i])
        job.translated = n
        if progress:
            progress(n, job.blocks)

    job.output = output_path(job.source, job.target_code)
    job.output.write_text("\n".join(lines) + "\n", encoding="utf-8")


# ----------------------------------------------------------------------- docx
def _docx_paragraph_texts(doc):
    """Every paragraph in the document, including tables, headers and footers."""
    seen = []

    def walk_paragraphs(container):
        for p in container.paragraphs:
            seen.append(p)
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    walk_paragraphs(cell)

    walk_paragraphs(doc)
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is not None:
                walk_paragraphs(part)
    return seen


def _set_paragraph_text(paragraph, text: str) -> None:
    """Put text into a paragraph, keeping its style and its first run's look."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _translate_docx(job: Job, translate, progress=None) -> None:
    import docx

    doc = docx.Document(str(job.source))
    paragraphs = [p for p in _docx_paragraph_texts(doc) if p.text.strip()]
    job.blocks = len(paragraphs)

    mixed = 0
    for n, p in enumerate(paragraphs, start=1):
        if len([r for r in p.runs if r.text.strip()]) > 1:
            mixed += 1
        _set_paragraph_text(p, translate(p.text))
        job.translated = n
        if progress:
            progress(n, job.blocks)

    if mixed:
        job.warnings.append(
            f"{mixed} paragraphs had formatting that changed part way through, such as a bold "
            "word inside a sentence. The paragraph's own style is kept, but that inline "
            "formatting is not, because translation moves the words."
        )

    job.output = output_path(job.source, job.target_code)
    doc.save(str(job.output))


# ------------------------------------------------------------------------ pdf
def _translate_pdf(job: Job, translate, progress=None) -> None:
    """PDF in, Word out. Said plainly rather than pretending to rebuild a PDF."""
    import docx
    from pypdf import PdfReader

    reader = PdfReader(str(job.source))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise RuntimeError("This PDF is password protected. Remove the password and try again.")

    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    if not any(pages):
        raise RuntimeError(
            "No text could be read from this PDF. It is probably a scan, which needs to be put "
            "through OCR before it can be translated."
        )

    blocks: list[tuple[int, str]] = []
    for i, text in enumerate(pages):
        paras = [p for p in re.split(r"\n\s*\n", text) if p.strip()]
        # Plenty of PDFs carry no blank lines at all, and extracting one of those
        # as a single block gives the model a page-long sentence to chew on. When
        # a page looks like that, treat each line as its own block instead.
        if len(paras) <= 1 and text.count("\n") > 2:
            paras = [ln for ln in text.split("\n") if ln.strip()]
        for para in paras:
            blocks.append((i, re.sub(r"\s*\n\s*", " ", para).strip()))
    job.blocks = len(blocks)

    out = docx.Document()
    current = -1
    for n, (page_index, para) in enumerate(blocks, start=1):
        if page_index != current and current != -1:
            out.add_page_break()
        current = page_index
        out.add_paragraph(translate(para))
        job.translated = n
        if progress:
            progress(n, job.blocks)

    job.warnings.append(
        "A PDF's layout cannot be rebuilt around translated text, so this was written as a Word "
        "document with the text in reading order and a page break between pages."
    )
    job.note = "PDF read as text, written as .docx"
    job.output = output_path(job.source, job.target_code, ".docx")
    out.save(str(job.output))


# --------------------------------------------------------------------- router
_HANDLERS = {
    ".txt": _translate_plain,
    ".md": _translate_plain,
    ".srt": _translate_subtitles,
    ".vtt": _translate_subtitles,
    ".docx": _translate_docx,
    ".pdf": _translate_pdf,
}


def translate_file(source: Path, target_code: str, translate, progress=None) -> Job:
    job = Job(source=source, target_code=target_code)
    ext = source.suffix.lower()
    handler = _HANDLERS.get(ext)
    if handler is None:
        raise RuntimeError(f"LingoGeek does not read {ext or 'files without an extension'} yet.")
    handler(job, translate, progress)
    return job
